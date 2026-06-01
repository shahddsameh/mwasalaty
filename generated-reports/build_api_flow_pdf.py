from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "mwaslaty-api-flow-support-review.docx"


BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "DFF3E6"
AMBER = "FFF2CC"
RED = "FCE4E4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="DADCE0", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = level <= 2
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, header, bold=True, color=BLUE, size=8.5)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_text(cell, str(value), size=8.2)
    doc.add_paragraph()
    return table


def add_status_table(doc):
    rows = [
        ("Flow 1: Plan A to B", "Partial", "POST /api/plan supports coordinates, date/time, modes, route search, route results, and no-route handling. Missing geocoding/autocomplete and AI prompt parsing."),
        ("Flow 2: Compare route options", "Partial", "Route cards include duration, estimated fare, transfers, modes, summary, and legs. Sorting exists for quickest, cheapest, most_comfortable, and reliable. Missing formal comfort score and route detail endpoint."),
        ("Flow 3: Follow step-by-step", "Mostly missing", "Legs include basic instructions, times, distances, and stops. Missing navigation sessions, GPS updates, deviation detection, recalculation, offline sync, arrival confirmation, and feedback."),
        ("Flow 4: Save places/routes", "Client-only", "The diagram stores data in IndexedDB, so no backend is required for local-only saving. Server sync endpoints are missing if account-based saved screens are needed."),
        ("Flow 5: AI Trip Planner", "Missing", "No endpoint for AI prompt analysis, budget/duration parsing, attractions/restaurants, multi-day schedules, cost breakdown, save, or share."),
        ("Flow 6: Booking and payment", "Missing", "No booking, quote, payment, ticket generation, QR payload, validation status, ticket scan, refund, or operator authorization API."),
    ]
    table = add_table(
        doc,
        ("Flow", "Support", "Evidence and gaps"),
        rows,
        (2200, 1300, 5860),
    )
    for row in table.rows[1:]:
        status = row.cells[1].text
        fill = GREEN if status == "Client-only" else AMBER if "Partial" in status or "Mostly" in status else RED
        set_cell_shading(row.cells[1], fill)


def add_edge_case_table(doc):
    rows = [
        ("No results", "404", "OTP_EMPTY_PLAN", "Return error details with from, to, date, and time. Current API already does this for empty itineraries."),
        ("Partial routes", "200 or 409", "ROUTE_PARTIAL_ONLY", "Return status partial with available legs, unavailable legs, and whether booking/navigation can continue."),
        ("Fare unknown", "200", "FARE_UNKNOWN warning", "Return fare amount null, currency EGP, status unknown, and fareConfidence 0 rather than silently using 0."),
        ("Off-hours", "422", "ROUTE_OUTSIDE_SERVICE_HOURS", "Return service window, timezone Africa/Cairo, and next available service time."),
        ("Signal loss mid-trip", "200/503", "NAVIGATION_LOCATION_UNAVAILABLE", "Keep session in offline_tracking when cached route exists. If server action is required, return recoverable unavailable error."),
        ("Ticket scanned twice", "409", "TICKET_ALREADY_USED", "Return original scannedAt, stationId, legId, and ticket status used."),
        ("Wrong-leg scan", "409", "TICKET_WRONG_LEG", "Return expected leg/route and scanned leg/route. Do not validate the ticket."),
        ("Expired ticket", "410", "TICKET_EXPIRED", "Return expiresAt and optional rebook action metadata."),
        ("Scan while offline", "202 or 503", "TICKET_SCAN_OFFLINE_PENDING", "Accept signed offline QR scans as pending sync. Otherwise require online validation."),
        ("Operator not authorized", "403", "OPERATOR_UNAUTHORIZED", "Reject scan request without leaking ticket details. Include required operator scope."),
    ]
    add_table(
        doc,
        ("Edge case", "HTTP", "Code", "Expected API response"),
        rows,
        (1900, 900, 2200, 4360),
        header_fill=LIGHT_GRAY,
    )


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "434343"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(4)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mwaslaty API Flow Support Review")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Backend support assessment for the six product flows and key edge cases")
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "The current backend exposes one route-planning endpoint, POST /api/plan. It is a strong first slice for route search and comparison, but it is not yet a full product API for navigation, AI trip planning, saved account data, payments, tickets, scanning, or operator workflows."
    )
    add_body(
        doc,
        "Existing support includes route search, route cards, estimated fares, transfers, modes, leg instructions, and basic OTP-related errors. Most remaining flows require new backend contracts and domain-specific error codes."
    )

    add_heading(doc, "Current API Surface", 1)
    add_bullets(doc, [
        "POST /api/plan validates from.lat, from.lng, to.lat, to.lng, date, and time.",
        "Optional preferences include modes and optimizeFor.",
        "Successful responses include planId, source, optimizedFor, from, to, itineraries, legs, durationMinutes, totalFare, transfers, summary, route names, and basic instructions.",
        "Current error codes: VALIDATION_ERROR, OTP_EMPTY_PLAN, OTP_SERVICE_UNAVAILABLE, OTP_GRAPHQL_ERROR, and INTERNAL_SERVER_ERROR.",
    ])

    add_heading(doc, "Six Flow Support Matrix", 1)
    add_status_table(doc)

    add_heading(doc, "Missing Endpoints", 1)
    add_bullets(doc, [
        "GET /api/places/search for destination suggestions and autocomplete.",
        "POST /api/ai/route-intent for natural-language route requests.",
        "GET /api/plans/:planId or GET /api/routes/:itineraryId for route detail reload.",
        "POST /api/navigation/sessions, PATCH /api/navigation/sessions/:id/location, and POST /api/navigation/sessions/:id/recalculate.",
        "POST /api/feedback.",
        "Saved data endpoints if account sync is required: /api/saved/places and /api/saved/routes.",
        "POST /api/ai/trips for the AI Trip Planner.",
        "POST /api/bookings/quote, POST /api/bookings, POST /api/payments/confirm.",
        "GET /api/tickets/:ticketId and POST /api/tickets/:ticketId/scan.",
        "Authentication and operator authorization endpoints.",
    ])

    add_heading(doc, "Missing Fields", 1)
    add_bullets(doc, [
        "Route geometry or polyline for map rendering.",
        "Stop IDs, GTFS IDs, platform names, and coordinates per leg.",
        "Detailed step list per leg beyond a single instruction string.",
        "fare.status, fare.source, and fareConfidence.",
        "comfortScore, crowding, disruption alerts, realtime delay, and accessibility metadata.",
        "Ticket eligibility, bookingId, paymentId, ticketId, qrPayload, expiresAt, and validationStatus.",
        "Offline/cache metadata such as cacheable, validUntil, and lastSyncedAt.",
    ])

    add_heading(doc, "Missing Error Codes", 1)
    add_bullets(doc, [
        "GEOCODING_NO_MATCH, AI_PARSE_FAILED, ROUTE_OUTSIDE_SERVICE_HOURS, ROUTE_PARTIAL_ONLY, FARE_UNKNOWN.",
        "NAVIGATION_OFFLINE_NO_CACHE, NAVIGATION_LOCATION_UNAVAILABLE, ROUTE_DEVIATION_RECALCULATION_FAILED.",
        "PAYMENT_FAILED, TICKET_EXPIRED, TICKET_ALREADY_USED, TICKET_WRONG_LEG, TICKET_SCAN_OFFLINE_PENDING.",
        "AUTH_REQUIRED, FORBIDDEN, OPERATOR_UNAUTHORIZED.",
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Edge Case Response Matrix", 1)
    add_edge_case_table(doc)

    add_heading(doc, "Recommended Next API Slice", 1)
    add_body(doc, "The most useful next backend increment is to harden route planning before starting payment or AI workflows.")
    add_bullets(doc, [
        "Add place search/autocomplete and route detail reload.",
        "Add route geometry, stop IDs, fare status, and comfort/reliability metadata.",
        "Separate no-route, off-hours, partial-route, and fare-unknown responses.",
        "Then add navigation sessions and feedback.",
        "Treat booking, payment, tickets, and operator scanning as a later bounded module with its own security model.",
    ])

    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Mwaslaty API review")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("777777")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
